"""Word document parser implementation."""

from pathlib import Path
from typing import Union
from docx import Document

from resume_parser.parsers.base import FileParser
from resume_parser.exceptions import FileParsingError
from resume_parser.utils.logger import setup_logger

logger = setup_logger(__name__)


class WordParser(FileParser):
    """
    Parse Word (.docx) resume files to extract text content.
    
    Uses python-docx to extract text from paragraphs, tables,
    and other document elements.
    """
    
    def parse(self, file_path: Union[str, Path]) -> str:
        """
        Extract text from a Word document.
        
        Args:
            file_path: Path to .docx file
        
        Returns:
            Extracted text content
        
        Raises:
            FileParsingError: If document cannot be parsed
            FileNotFoundError: If file doesn't exist
        """
        # Validate file exists
        path = self._validate_file_exists(file_path)
        self._validate_file_size(path)
        
        logger.info(f"Parsing Word document: {path}")
        
        try:
            doc = Document(path)
            text_content = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            full_text = "\n".join(text_content)
            
            if not full_text.strip():
                raise FileParsingError(
                    f"No text content extracted from Word document: {path}"
                )
            
            logger.info(
                f"Successfully extracted {len(full_text)} characters from Word document"
            )
            return full_text
            
        except FileParsingError:
            raise
        except Exception as e:
            logger.error(f"Unexpected error parsing Word document {path}: {str(e)}")
            raise FileParsingError(f"Failed to parse Word document: {str(e)}")

